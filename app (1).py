import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

st.set_page_config(page_title="Generator Modul Ajar - Kurikulum Merdeka", layout="centered")

st.title("📘 Generator Modul Ajar Profesional")
st.caption("Sesuai Kurikulum Merdeka — Pendekatan Pembelajaran Mendalam (Deep Learning): Berkesadaran, Bermakna, Menggembirakan")

# =========================================================
# FORM INPUT
# =========================================================
with st.form("input_form"):

    st.subheader("A. Informasi Umum")

    col1, col2 = st.columns(2)
    with col1:
        nama_guru = st.text_input("Nama Penyusun")
        institusi = st.text_input("Nama Sekolah / Institusi")
        tahun = st.text_input("Tahun Penyusunan", value="2026")
    with col2:
        jenjang = st.selectbox("Jenjang", ["SD/MI", "SMP/MTs", "SMA/MA", "SMK"])
        kelas = st.text_input("Kelas", value="VII")
        semester = st.selectbox("Semester", ["Ganjil", "Genap"])

    col3, col4 = st.columns(2)
    with col3:
        mapel = st.selectbox("Mata Pelajaran", ["Matematika", "Bahasa Indonesia", "IPA", "IPS", "Bahasa Inggris", "PJOK", "Seni Budaya", "Lainnya"])
        fase = st.selectbox("Fase", ["Fase A", "Fase B", "Fase C", "Fase D", "Fase E", "Fase F"])
    with col4:
        alokasi_waktu = st.text_input("Alokasi Waktu (menit)", value="80")
        jumlah_pertemuan = st.text_input("Jumlah Pertemuan", value="1")

    bab_materi = st.text_input("Bab / Materi Pokok")

    kompetensi_awal = st.text_area("Kompetensi Awal (kemampuan prasyarat peserta didik)")

    profil_pelajar = st.multiselect(
        "Profil Pelajar Pancasila",
        ["Beriman, Bertakwa kepada Tuhan YME, dan Berakhlak Mulia", "Berkebinekaan Global",
         "Bergotong Royong", "Mandiri", "Bernalar Kritis", "Kreatif"],
        default=["Bernalar Kritis", "Kreatif"]
    )

    sarana_prasarana = st.text_area("Sarana dan Prasarana", value="Papan tulis, spidol, LKPD, laptop, proyektor, koneksi internet")

    target_peserta_didik = st.selectbox(
        "Target Peserta Didik",
        ["Peserta didik reguler/tipikal", "Peserta didik dengan kesulitan belajar",
         "Peserta didik dengan pencapaian tinggi", "Peserta didik cerdas istimewa berbakat istimewa (CIBI)"]
    )

    model_pembelajaran = st.selectbox(
        "Model Pembelajaran",
        ["Discovery Learning", "Problem Based Learning (PBL)", "Project Based Learning (PjBL)",
         "Inquiry Learning", "Cooperative Learning", "Tatap Muka", "Blended Learning"]
    )

    st.divider()
    st.subheader("B. Komponen Inti")

    tp = st.text_area("Tujuan Pembelajaran (pisahkan tiap poin dengan koma)", height=120)
    pemahaman_bermakna = st.text_area("Pemahaman Bermakna",
        help="Manfaat/relevansi materi ini bagi kehidupan peserta didik")
    pertanyaan_pemantik = st.text_area("Pertanyaan Pemantik (pisahkan tiap pertanyaan dengan koma)")

    st.markdown("**Kegiatan Pembelajaran**")
    kegiatan_pendahuluan = st.text_area("Kegiatan Pendahuluan",
        value="Guru membuka pembelajaran dengan salam dan doa, memeriksa kehadiran, menyampaikan tujuan pembelajaran, dan memberikan apersepsi.")

    st.markdown("*Kegiatan Inti — Pendekatan Pembelajaran Mendalam (Deep Learning)*")
    kegiatan_berkesadaran = st.text_area("🧘 Berkesadaran (Mindful Learning)",
        help="Aktivitas yang membangun kesadaran diri, fokus, dan niat belajar peserta didik",
        value="Peserta didik diajak merefleksikan tujuan belajar hari ini dan menghubungkannya dengan pengalaman pribadi.")
    kegiatan_bermakna = st.text_area("🎯 Bermakna (Meaningful Learning)",
        help="Aktivitas inti yang menghubungkan materi dengan konteks nyata dan membangun pemahaman mendalam",
        value="Peserta didik melakukan eksplorasi, diskusi kelompok, dan pemecahan masalah kontekstual terkait materi.")
    kegiatan_menggembirakan = st.text_area("😊 Menggembirakan (Joyful Learning)",
        help="Aktivitas yang membuat proses belajar menyenangkan dan memotivasi",
        value="Peserta didik melakukan permainan edukatif, presentasi kreatif, atau kuis interaktif terkait materi.")

    kegiatan_penutup = st.text_area("Kegiatan Penutup",
        value="Guru dan peserta didik menyimpulkan pembelajaran, melakukan refleksi, dan menyampaikan rencana pertemuan berikutnya.")

    st.markdown("**Asesmen**")
    asesmen_diagnostik = st.text_area("Asesmen Diagnostik (di awal pembelajaran)")
    asesmen_formatif = st.text_area("Asesmen Formatif (selama proses pembelajaran)")
    asesmen_sumatif = st.text_area("Asesmen Sumatif (di akhir pembelajaran/unit)")

    pengayaan = st.text_area("Pengayaan (untuk peserta didik yang telah mencapai tujuan pembelajaran)")
    remedial = st.text_area("Remedial (untuk peserta didik yang belum mencapai tujuan pembelajaran)")

    st.divider()
    st.subheader("C. Lampiran")
    lkpd = st.text_area("LKPD / Lembar Kerja Peserta Didik (opsional)")
    bahan_bacaan = st.text_area("Bahan Bacaan Guru dan Peserta Didik (opsional)")
    glosarium = st.text_area("Glosarium (istilah, pisahkan dengan koma) (opsional)")
    daftar_pustaka = st.text_area("Daftar Pustaka (opsional)")

    submitted = st.form_submit_button("Buat Modul Ajar", use_container_width=True)

# =========================================================
# FUNGSI BANTUAN
# =========================================================
def tambah_bullet(doc, teks, pemisah=','):
    for item in teks.split(pemisah):
        item = item.strip()
        if item:
            doc.add_paragraph(item, style='List Bullet')

# =========================================================
# LOGIKA PEMROSESAN
# =========================================================
if submitted:
    if not nama_guru.strip() or not tp.strip():
        st.error("Mohon isi minimal Nama Penyusun dan Tujuan Pembelajaran.")
    else:
        doc = Document()

        # Judul
        judul = doc.add_heading('MODUL AJAR', level=0)
        judul.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # A. INFORMASI UMUM
        doc.add_heading('A. Informasi Umum', level=1)
        tabel = doc.add_table(rows=0, cols=2)
        tabel.style = 'Table Grid'
        data_identitas = [
            ("Nama Penyusun", nama_guru),
            ("Institusi", institusi),
            ("Tahun Penyusunan", tahun),
            ("Jenjang", jenjang),
            ("Kelas / Semester", f"{kelas} / {semester}"),
            ("Mata Pelajaran", mapel),
            ("Bab / Materi Pokok", bab_materi),
            ("Fase", fase),
            ("Alokasi Waktu", f"{alokasi_waktu} menit x {jumlah_pertemuan} pertemuan"),
            ("Model Pembelajaran", model_pembelajaran),
            ("Target Peserta Didik", target_peserta_didik),
        ]
        for label, isi in data_identitas:
            row = tabel.add_row().cells
            row[0].text = label
            row[1].text = str(isi) if isi else "-"

        doc.add_paragraph()
        doc.add_heading('Kompetensi Awal', level=2)
        doc.add_paragraph(kompetensi_awal.strip() if kompetensi_awal.strip() else "-")

        doc.add_heading('Profil Pelajar Pancasila', level=2)
        if profil_pelajar:
            for item in profil_pelajar:
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph("-")

        doc.add_heading('Sarana dan Prasarana', level=2)
        doc.add_paragraph(sarana_prasarana.strip() if sarana_prasarana.strip() else "-")

        # B. KOMPONEN INTI
        doc.add_heading('B. Komponen Inti', level=1)

        doc.add_heading('Tujuan Pembelajaran', level=2)
        tambah_bullet(doc, tp)

        doc.add_heading('Pemahaman Bermakna', level=2)
        doc.add_paragraph(pemahaman_bermakna.strip() if pemahaman_bermakna.strip() else "-")

        doc.add_heading('Pertanyaan Pemantik', level=2)
        if pertanyaan_pemantik.strip():
            tambah_bullet(doc, pertanyaan_pemantik)
        else:
            doc.add_paragraph("-")

        doc.add_heading('Kegiatan Pembelajaran', level=2)

        doc.add_heading('1. Pendahuluan', level=3)
        doc.add_paragraph(kegiatan_pendahuluan.strip() if kegiatan_pendahuluan.strip() else "-")

        doc.add_heading('2. Kegiatan Inti (Pendekatan Pembelajaran Mendalam)', level=3)
        doc.add_heading('a. Berkesadaran (Mindful Learning)', level=4)
        doc.add_paragraph(kegiatan_berkesadaran.strip() if kegiatan_berkesadaran.strip() else "-")
        doc.add_heading('b. Bermakna (Meaningful Learning)', level=4)
        doc.add_paragraph(kegiatan_bermakna.strip() if kegiatan_bermakna.strip() else "-")
        doc.add_heading('c. Menggembirakan (Joyful Learning)', level=4)
        doc.add_paragraph(kegiatan_menggembirakan.strip() if kegiatan_menggembirakan.strip() else "-")

        doc.add_heading('3. Penutup', level=3)
        doc.add_paragraph(kegiatan_penutup.strip() if kegiatan_penutup.strip() else "-")

        doc.add_heading('Asesmen', level=2)
        doc.add_heading('Asesmen Diagnostik', level=3)
        doc.add_paragraph(asesmen_diagnostik.strip() if asesmen_diagnostik.strip() else "-")
        doc.add_heading('Asesmen Formatif', level=3)
        doc.add_paragraph(asesmen_formatif.strip() if asesmen_formatif.strip() else "-")
        doc.add_heading('Asesmen Sumatif', level=3)
        doc.add_paragraph(asesmen_sumatif.strip() if asesmen_sumatif.strip() else "-")

        doc.add_heading('Pengayaan dan Remedial', level=2)
        doc.add_heading('Pengayaan', level=3)
        doc.add_paragraph(pengayaan.strip() if pengayaan.strip() else "-")
        doc.add_heading('Remedial', level=3)
        doc.add_paragraph(remedial.strip() if remedial.strip() else "-")

        # C. LAMPIRAN
        doc.add_heading('C. Lampiran', level=1)

        doc.add_heading('LKPD (Lembar Kerja Peserta Didik)', level=2)
        doc.add_paragraph(lkpd.strip() if lkpd.strip() else "-")

        doc.add_heading('Bahan Bacaan Guru dan Peserta Didik', level=2)
        doc.add_paragraph(bahan_bacaan.strip() if bahan_bacaan.strip() else "-")

        doc.add_heading('Glosarium', level=2)
        if glosarium.strip():
            tambah_bullet(doc, glosarium)
        else:
            doc.add_paragraph("-")

        doc.add_heading('Daftar Pustaka', level=2)
        doc.add_paragraph(daftar_pustaka.strip() if daftar_pustaka.strip() else "-")

        # Simpan ke buffer
        bio = io.BytesIO()
        doc.save(bio)

        st.success("✅ Modul Ajar lengkap berhasil dibuat!")
        st.download_button(
            label="⬇️ Download Modul Ajar (.docx)",
            data=bio.getvalue(),
            file_name=f"Modul_Ajar_{mapel}_{kelas}_{fase}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
